"""Document text extraction: PDF, DOCX, and URL HTML strip."""

from __future__ import annotations

import io
import re
from typing import Optional

import httpx
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return _normalize_whitespace("\n\n".join(parts))


def extract_text_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Tables often hold skills sections
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _normalize_whitespace("\n".join(parts))


def extract_text_from_upload(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if lower.endswith(".docx"):
        return extract_text_from_docx(data)
    if lower.endswith(".doc"):
        raise ValueError("Legacy .doc is not supported; please upload PDF or DOCX.")
    if lower.endswith(".txt") or lower.endswith(".md"):
        return _normalize_whitespace(data.decode("utf-8", errors="replace"))
    # Sniff PDF magic
    if data[:4] == b"%PDF":
        return extract_text_from_pdf(data)
    raise ValueError(f"Unsupported file type: {filename}. Use PDF or DOCX.")


def fetch_url_text(url: str, timeout: float = 30.0) -> str:
    headers = {"User-Agent": "ResumeMatch/0.1 (+local; job-description fetcher)"}
    with httpx.Client(timeout=timeout, follow_redirects=True, headers=headers) as client:
        resp = client.get(url)
        resp.raise_for_status()
        content_type = resp.headers.get("content-type", "")
        if "pdf" in content_type or url.lower().endswith(".pdf"):
            return extract_text_from_pdf(resp.content)
        html = resp.text
    return html_to_text(html)


def html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "nav", "footer", "header", "aside"]):
        tag.decompose()
    # Prefer main / article if present
    root = soup.find("main") or soup.find("article") or soup.body or soup
    text = root.get_text(separator="\n")
    return _normalize_whitespace(text)


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def resolve_jd_text(
    *,
    jd_text: Optional[str] = None,
    jd_url: Optional[str] = None,
    jd_filename: Optional[str] = None,
    jd_bytes: Optional[bytes] = None,
) -> str:
    if jd_text and jd_text.strip():
        return _normalize_whitespace(jd_text)
    if jd_url and jd_url.strip():
        return fetch_url_text(jd_url.strip())
    if jd_bytes is not None and jd_filename:
        return extract_text_from_upload(jd_filename, jd_bytes)
    raise ValueError("Provide jd_text, jd_url, or a JD file upload.")
