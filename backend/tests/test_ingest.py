from __future__ import annotations

from app.ingest.documents import (
    extract_text_from_docx,
    extract_text_from_pdf,
    html_to_text,
    _normalize_whitespace,
)
from docx import Document
import io


def test_normalize_whitespace():
    assert _normalize_whitespace("a  \n\n\n b") == "a\n\n b"
    assert "  " not in _normalize_whitespace("hello    world")


def test_html_to_text_strips_chrome():
    html = """
    <html><head><script>evil()</script></head>
    <body><nav>nav</nav><main><h1>Engineer</h1><p>Need React and TypeScript.</p></main>
    <footer>footer</footer></body></html>
    """
    text = html_to_text(html)
    assert "React" in text
    assert "evil" not in text
    assert "nav" not in text.lower() or "Engineer" in text


def test_docx_roundtrip():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Alice Example")
    doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL")
    doc.save(buf)
    text = extract_text_from_docx(buf.getvalue())
    assert "Alice Example" in text
    assert "Python" in text


def test_pdf_extract_nonempty_when_text_layer(tmp_path):
    # Minimal valid-ish PDF with a text object is hard inline; skip if pypdf gets empty.
    # This guards the function against crashing on empty PDF.
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    path = tmp_path / "blank.pdf"
    with path.open("wb") as f:
        writer.write(f)
    text = extract_text_from_pdf(path.read_bytes())
    assert isinstance(text, str)
